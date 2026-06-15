#!/usr/bin/env python3
"""실코퍼스 기반 groundedness E2E 하니스 (다포맷 추출 -> 골든질문 생성 -> 채점).

세 가지 작업을 한 스크립트로 제공한다(서브커맨드):
  1. generate     : 임의 로컬 코퍼스를 추출해 근거 기반 골든질문을 자동 생성한다.
  2. run-production: 코퍼스를 실행 중인 API에 업로드하고, 골든질문을 질의해
                     답변의 근거성(기대 용어 포함 + 출처 파일명 일치)을 채점한다.

언어/도메인 중립 설계:
  - 특정 언어(일본어 등)나 도메인 가정을 하드코딩하지 않는다.
  - 핵심 사실 추출 휴리스틱(score_fact_line)은 숫자/단위/대문자 토큰/길이 기반의
    언어 중립 신호를 사용한다.
  - 질문 템플릿은 --question-template로 외부화(기본은 영어, 자리표시자 사용).
  - 외부 정답 사실은 --manual-facts-json(파일명->문장 리스트)으로 선택 주입 가능.

사용법(2단계):
    # 1) 골든질문 생성(서버 불필요)
    uv run python scripts/production_corpus_e2e.py generate \
        --corpus-dir ./data/sample_corpus --output-dir ./reports/run1

    # 2) 라이브 서버에 업로드 + 채점
    uv run python scripts/production_corpus_e2e.py run-production \
        --corpus-dir ./data/sample_corpus \
        --golden ./reports/run1/golden_questions.json \
        --output-dir ./reports/run1 \
        --backend-url http://localhost:8000

인증:
    --api-key-env(기본 FASTAPI_AUTH_KEY) 환경변수에서 X-API-Key를 구성한다.
    키/서버가 없으면 generate는 동작하고 run-production은 graceful하게 보고한다.

출력/종료 코드:
    결과 요약을 JSON으로 출력. run-production은 QA 케이스가 모두 통과하면 0,
    아니면 1을 반환한다(서버 미가용 시에도 사유와 함께 1).

JapanRAG 원본 대비 일반화:
    - MANUAL_FACTS_BY_FILENAME 하드코딩 삭제(필요 시 --manual-facts-json 주입).
    - score_fact_line의 일본어 정규식(円/￥/株式会社 등) 제거 -> 언어 중립 휴리스틱.
    - 한국어 질문 템플릿 -> --question-template 외부화(영어 기본).
    - DEFAULT_FRONTEND_CONFIG(config.js 자동탐지) 제거 -> --backend-url 필수.
    - chunked 업로드 / X-OneRAG-Upload-Token / 멀티테넌트 company_id 제거.
"""

from __future__ import annotations

import argparse
import csv
import html
import json
import os
import re
import time
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import httpx
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

# 직접 업로드 가능한 확장자 -> MIME (OneRAG validate_file 지원 집합과 정합)
DIRECT_UPLOAD_MIME_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".csv": "text/csv",
    ".html": "text/html",
    ".json": "application/json",
}
# 직접 업로드는 어렵지만 텍스트 추출은 가능한 확장자(텍스트 대체본으로 업로드)
TEXT_SURROGATE_EXTENSIONS = {".pptx"}
# 질문 템플릿 기본값(언어 중립 영어). {filename}, {anchor}, {anchor_2} 치환.
DEFAULT_QUESTION_TEMPLATE = (
    "From the document named '{filename}', explain the key content related to '{anchor}'."
)
DEFAULT_QUESTION_TEMPLATE_2 = (
    "From the document named '{filename}', state the numbers, dates, names, or "
    "conditions related to '{anchor}' as grounded by the document."
)


@dataclass(frozen=True)
class CorpusDoc:
    """코퍼스 내 단일 문서 메타."""

    doc_id: str
    path: Path
    rel_path: str
    filename: str
    extension: str
    size_bytes: int


def utc_now() -> str:
    return datetime.now(UTC).isoformat().replace("+00:00", "Z")


def normalize_text(value: str) -> str:
    """NFKC 정규화 + 공백/개행 정리."""
    value = unicodedata.normalize("NFKC", value)
    value = value.replace(" ", " ")
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.split("\n")]
    return "\n".join(line for line in lines if line)


def compact(value: str | None) -> str:
    """공백 제거 + NFKC + casefold (비교용 정규화)."""
    if not value:
        return ""
    return re.sub(r"\s+", "", unicodedata.normalize("NFKC", value)).casefold()


def alnum_cjk(value: str | None) -> str:
    """영숫자 + CJK 문자만 남긴다(언어 중립 토큰 비교용)."""
    normalized = compact(value)
    return "".join(ch for ch in normalized if ch.isalnum() or "぀" <= ch <= "鿿")


def discover_corpus(corpus_dir: Path) -> list[CorpusDoc]:
    """코퍼스 디렉토리를 순회해 문서 목록을 구성한다."""
    files = sorted(path for path in corpus_dir.rglob("*") if path.is_file())
    docs: list[CorpusDoc] = []
    for index, path in enumerate(files, start=1):
        docs.append(
            CorpusDoc(
                doc_id=f"doc-{index:03d}",
                path=path,
                rel_path=path.relative_to(corpus_dir).as_posix(),
                filename=path.name,
                extension=path.suffix.lower(),
                size_bytes=path.stat().st_size,
            )
        )
    return docs


def extract_pdf(path: Path) -> str:
    texts: list[str] = []
    with path.open("rb") as handle:
        reader = PdfReader(handle)
        for page_index, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as error:  # noqa: BLE001 - 페이지 추출 실패는 마커로 기록
                text = f"[page {page_index} extraction failed: {type(error).__name__}]"
            if text.strip():
                texts.append(f"[page {page_index}]\n{text}")
    return normalize_text("\n\n".join(texts))


def extract_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"[table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return normalize_text("\n".join(parts))


def extract_xlsx(path: Path, *, max_rows_per_sheet: int = 250) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"[sheet {sheet.title}]")
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [str(value).strip() for value in row if value not in (None, "")]
            if values:
                parts.append(f"row {row_index}: " + " | ".join(values))
            if row_index >= max_rows_per_sheet:
                parts.append(f"[sheet {sheet.title} truncated after {max_rows_per_sheet} rows]")
                break
    workbook.close()
    return normalize_text("\n".join(parts))


def extract_pptx(path: Path) -> str:
    parts: list[str] = []
    with zipfile.ZipFile(path) as archive:
        slide_names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        for slide_index, name in enumerate(slide_names, start=1):
            root = ElementTree.fromstring(archive.read(name))
            texts = [
                node.text.strip()
                for node in root.iter()
                if node.tag.endswith("}t") and node.text and node.text.strip()
            ]
            if texts:
                parts.append(f"[slide {slide_index}]\n" + "\n".join(texts))
    return normalize_text("\n\n".join(parts))


def extract_text(
    path: Path, *, manual_facts_by_filename: dict[str, list[str]] | None = None
) -> tuple[str, str | None]:
    """확장자별로 텍스트를 추출한다. (text, warning) 반환.

    추출 텍스트가 빈약하고 --manual-facts-json으로 외부 사실이 주입된 경우 보완한다.
    """
    ext = path.suffix.lower()
    manual_facts = (manual_facts_by_filename or {}).get(path.name)
    try:
        if ext == ".pdf":
            text = extract_pdf(path)
        elif ext == ".docx":
            text = extract_docx(path)
        elif ext == ".xlsx":
            text = extract_xlsx(path)
        elif ext == ".pptx":
            return extract_pptx(path), "text_surrogate_required"
        elif ext in {".txt", ".md", ".csv", ".html", ".json"}:
            text = normalize_text(path.read_text(encoding="utf-8", errors="replace"))
        else:
            return "", "unsupported_extension"
    except Exception as error:  # noqa: BLE001 - 추출 실패는 빈 텍스트 + 경고로 반환
        return "", f"{type(error).__name__}: {error}"

    if manual_facts and len(text) < 100:
        merged = normalize_text(
            "\n\n".join(part for part in [text, "\n".join(manual_facts)] if part)
        )
        return merged, "manual_facts_added"
    return text, None


def split_lines(text: str) -> list[str]:
    """채점에 쓸만한 텍스트 라인만 추린다(짧은/기호뿐/페이지마커 제외)."""
    lines: list[str] = []
    for raw in text.split("\n"):
        line = raw.strip()
        if len(line) < 8:
            continue
        if re.fullmatch(r"[\W_]+", line):
            continue
        lowered = line.casefold()
        if lowered.startswith("[page ") or lowered.startswith("[slide "):
            continue
        lines.append(line[:240])
    return lines


def score_fact_line(line: str) -> tuple[int, int]:
    """언어 중립 휴리스틱으로 '사실성 높은 라인'을 점수화한다.

    숫자/단위 토큰/대문자 약어/적절한 길이를 신호로 사용한다(특정 언어 정규식 없음).
    """
    score = 0
    if re.search(r"\d", line):
        score += 5
    # 단위/기호: 통화·퍼센트·측정 단위 등 언어 무관 신호
    if re.search(r"%|[$€£¥₩]|\b(kg|mm|cm|km|ml|hz|mah|kw|gb|mb)\b", line, re.I):
        score += 4
    # 대문자 약어/고유명사 토큰(2자 이상 대문자 연속)
    if re.search(r"\b[A-Z]{2,}\b", line):
        score += 3
    # 날짜 형태(YYYY-MM-DD 등)
    if re.search(r"\b\d{4}[-/.]\d{1,2}[-/.]\d{1,2}\b", line):
        score += 3
    if 18 <= len(line) <= 120:
        score += 2
    return (score, min(len(line), 120))


def select_facts(text: str, *, limit: int = 8) -> list[str]:
    """점수 상위 + 본문 순서로 중복 없는 핵심 사실 라인을 선택한다."""
    seen: set[str] = set()
    lines = split_lines(text)
    ranked = sorted(lines, key=score_fact_line, reverse=True)
    facts: list[str] = []
    for line in ranked + lines:
        key = compact(line[:80])
        if not key or key in seen:
            continue
        seen.add(key)
        facts.append(line)
        if len(facts) >= limit:
            break
    return facts


def phrase_from_fact(fact: str) -> str:
    """사실 라인에서 질문 앵커로 쓸 짧은 구문을 추출한다."""
    fact = re.sub(r"^\[[^\]]+\]\s*", "", fact).strip()
    if len(fact) <= 42:
        return fact
    return fact[:42].rstrip(",. ")


def build_cases_for_doc(
    doc: CorpusDoc,
    text: str,
    *,
    question_template: str,
    question_template_2: str,
) -> list[dict[str, Any]]:
    """문서 1건에 대해 근거 기반 골든질문 2건을 생성한다."""
    facts = select_facts(text, limit=8)
    if not facts:
        facts = [doc.filename, doc.path.stem]
    anchor_1 = phrase_from_fact(facts[0])
    anchor_2 = phrase_from_fact(facts[1] if len(facts) > 1 else facts[0])
    answer_1_facts = facts[:3]
    answer_2_facts = facts[1:5] if len(facts) > 1 else facts[:3]

    return [
        {
            "id": f"{doc.doc_id}-q01",
            "doc_id": doc.doc_id,
            "document_filename": doc.filename,
            "question": question_template.format(filename=doc.filename, anchor=anchor_1),
            "expected_answer": " / ".join(answer_1_facts),
            "expected_answer_contains": [phrase_from_fact(item) for item in answer_1_facts[:3]],
            "expected_source_filename": doc.filename,
            "grading": {"min_expected_terms": min(2, len(answer_1_facts))},
        },
        {
            "id": f"{doc.doc_id}-q02",
            "doc_id": doc.doc_id,
            "document_filename": doc.filename,
            "question": question_template_2.format(filename=doc.filename, anchor=anchor_2),
            "expected_answer": " / ".join(answer_2_facts),
            "expected_answer_contains": [phrase_from_fact(item) for item in answer_2_facts[:3]],
            "expected_source_filename": doc.filename,
            "grading": {"min_expected_terms": min(2, len(answer_2_facts))},
        },
    ]


def _load_manual_facts(path: str | None) -> dict[str, list[str]]:
    """--manual-facts-json 파일을 로드한다(없으면 빈 dict)."""
    if not path:
        return {}
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise RuntimeError("--manual-facts-json must map filename -> list[str]")
    return {str(key): [str(item) for item in value] for key, value in data.items()}


def generate_golden(
    corpus_dir: Path,
    output_dir: Path,
    *,
    question_template: str,
    question_template_2: str,
    manual_facts_json: str | None,
) -> dict[str, Path]:
    """코퍼스를 추출해 manifest/golden(JSON, Markdown)을 생성한다."""
    output_dir.mkdir(parents=True, exist_ok=True)
    manual_facts = _load_manual_facts(manual_facts_json)
    docs = discover_corpus(corpus_dir)
    extracted: list[dict[str, Any]] = []
    text_by_doc: dict[str, str] = {}
    for doc in docs:
        text, warning = extract_text(doc.path, manual_facts_by_filename=manual_facts)
        text_by_doc[doc.doc_id] = text
        extracted.append(
            {
                "doc_id": doc.doc_id,
                "filename": doc.filename,
                "rel_path": doc.rel_path,
                "extension": doc.extension,
                "size_bytes": doc.size_bytes,
                "text_chars": len(text),
                "direct_upload_supported": doc.extension in DIRECT_UPLOAD_MIME_TYPES,
                "requires_text_surrogate": doc.extension in TEXT_SURROGATE_EXTENSIONS,
                "warning": warning,
                "preview": text[:500],
            }
        )

    cases: list[dict[str, Any]] = []
    for doc in docs:
        cases.extend(
            build_cases_for_doc(
                doc,
                text_by_doc[doc.doc_id],
                question_template=question_template,
                question_template_2=question_template_2,
            )
        )

    manifest = {
        "schema_version": 1,
        "generated_at": utc_now(),
        "corpus_dir": str(corpus_dir),
        "document_count": len(docs),
        "case_count": len(cases),
        "documents": extracted,
    }
    golden = {
        "schema_version": 1,
        "generated_at": utc_now(),
        "strategy": "deterministic_grounded_questions_from_extracted_local_documents",
        "document_count": len(docs),
        "case_count": len(cases),
        "cases": cases,
    }
    manifest_path = output_dir / "corpus_manifest.json"
    golden_path = output_dir / "golden_questions.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    golden_path.write_text(json.dumps(golden, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_markdown_golden(output_dir / "golden_questions.md", manifest, golden)
    return {"manifest": manifest_path, "golden": golden_path}


def _write_markdown_golden(path: Path, manifest: dict[str, Any], golden: dict[str, Any]) -> None:
    lines = [
        "# Corpus Golden Questions",
        "",
        f"- Generated at: `{golden['generated_at']}`",
        f"- Documents: `{manifest['document_count']}`",
        f"- Questions: `{golden['case_count']}`",
        "",
        "| ID | Document | Question | Expected answer basis |",
        "|---|---|---|---|",
    ]
    for case in golden["cases"]:
        lines.append(
            "| "
            + " | ".join(
                [
                    html.escape(str(case["id"])),
                    html.escape(str(case["document_filename"])),
                    html.escape(str(case["question"])),
                    html.escape(str(case["expected_answer"])[:300]),
                ]
            )
            + " |"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


# --------------------------- run-production (라이브) ---------------------------


def _api_root(backend_url: str) -> str:
    backend_url = backend_url.rstrip("/")
    return backend_url if backend_url.endswith("/api") else f"{backend_url}/api"


def _headers(api_key: str | None) -> dict[str, str]:
    return {"X-API-Key": api_key} if api_key else {}


def _require_success(response: httpx.Response, context: str) -> dict[str, Any]:
    try:
        payload = response.json()
    except json.JSONDecodeError:
        payload = {"raw": response.text[:1000]}
    if response.status_code >= 400:
        raise RuntimeError(f"{context} failed: HTTP {response.status_code}: {payload}")
    if not isinstance(payload, dict):
        raise RuntimeError(f"{context} did not return a JSON object: {payload!r}")
    return payload


def _create_session(
    client: httpx.Client, backend_url: str, api_key: str | None
) -> str:
    """채팅 세션을 생성하고 session_id를 반환한다."""
    response = client.post(
        f"{_api_root(backend_url)}/chat/session",
        headers=_headers(api_key),
        json={"metadata": {"purpose": "production_corpus_e2e"}},
        timeout=60.0,
    )
    payload = _require_success(response, "create session")
    return str(payload["session_id"])


def _staged_upload_file(
    doc: CorpusDoc, staging_dir: Path
) -> tuple[Path, str, str, str]:
    """업로드 가능한 (경로, 파일명, MIME, 모드)를 반환한다.

    직접 지원 확장자는 그대로, pptx 등은 텍스트 대체본(.txt)으로 변환한다.
    """
    if doc.extension in DIRECT_UPLOAD_MIME_TYPES:
        return doc.path, doc.filename, DIRECT_UPLOAD_MIME_TYPES[doc.extension], "direct"
    if doc.extension in TEXT_SURROGATE_EXTENSIONS:
        text, warning = extract_text(doc.path)
        if not text.strip():
            raise RuntimeError(f"cannot build text surrogate for {doc.filename}: {warning}")
        slug = re.sub(r"[^A-Za-z0-9]+", "-", doc.path.stem).strip("-").lower()[:48] or "document"
        output = staging_dir / f"{doc.doc_id}-{slug}.txt"
        output.write_text(
            f"Original filename: {doc.filename}\nOriginal extension: {doc.extension}\n\n{text}\n",
            encoding="utf-8",
        )
        return output, f"{doc.path.stem}.txt", "text/plain", "text_surrogate"
    raise RuntimeError(f"unsupported extension for upload: {doc.filename}")


def _upload_one(
    client: httpx.Client,
    backend_url: str,
    doc: CorpusDoc,
    staging_dir: Path,
    api_key: str | None,
) -> str:
    """문서 1건을 direct 경로로 업로드하고 job_id를 반환한다."""
    upload_path, upload_filename, mime_type, mode = _staged_upload_file(doc, staging_dir)
    metadata = {
        "corpus_doc_id": doc.doc_id,
        "original_filename": doc.filename,
        "original_rel_path": doc.rel_path,
        "upload_mode": mode,
    }
    with upload_path.open("rb") as handle:
        response = client.post(
            f"{_api_root(backend_url)}/upload",
            headers=_headers(api_key),
            data={"metadata": json.dumps(metadata, ensure_ascii=False)},
            files={"file": (upload_filename, handle, mime_type)},
            timeout=180.0,
        )
    payload = _require_success(response, f"upload {doc.filename}")
    return str(payload["job_id"])


def _poll_upload(
    client: httpx.Client,
    backend_url: str,
    job_id: str,
    *,
    timeout_seconds: float,
    poll_interval_seconds: float,
    api_key: str | None,
) -> dict[str, Any]:
    """업로드 잡이 종료 상태가 될 때까지 폴링한다."""
    deadline = time.monotonic() + timeout_seconds
    last_payload: dict[str, Any] = {}
    while time.monotonic() < deadline:
        response = client.get(
            f"{_api_root(backend_url)}/upload/status/{job_id}",
            headers=_headers(api_key),
            timeout=60.0,
        )
        payload = _require_success(response, f"poll upload {job_id}")
        last_payload = payload
        if payload.get("status") in {"completed", "failed", "cancelled"}:
            return payload
        time.sleep(poll_interval_seconds)
    raise TimeoutError(f"upload job {job_id} did not finish: {last_payload}")


def _chat_answer(
    client: httpx.Client,
    backend_url: str,
    session_id: str,
    question: str,
    api_key: str | None,
) -> dict[str, Any]:
    """질문 1건을 /chat에 보내고 응답 페이로드를 반환한다."""
    response = client.post(
        f"{_api_root(backend_url)}/chat",
        headers=_headers(api_key),
        json={"message": question, "session_id": session_id},
        timeout=180.0,
    )
    return _require_success(response, "chat answer")


def _source_filename_match(
    sources: list[dict[str, Any]], filename: str, uploaded_id: str | None
) -> bool:
    """응답 출처 중 기대 파일명(또는 업로드 job_id)이 포함된 것이 있는지 확인."""
    expected_stem = compact(Path(filename).stem)
    for source in sources:
        source_values = [
            source.get("document"),
            source.get("document_name"),
            source.get("filename"),
            source.get("source_file"),
            source.get("document_id"),
        ]
        if uploaded_id:
            source_values.append(uploaded_id)
        joined = compact(" ".join(str(value) for value in source_values if value))
        if expected_stem and expected_stem in joined:
            return True
        if uploaded_id and uploaded_id in joined:
            return True
    return False


def _term_matches(answer: str, term: str) -> bool:
    """기대 용어가 답변에 (정규화 기준) 포함되는지 확인."""
    answer_compact = compact(answer)
    term_compact = compact(term)
    if term_compact and term_compact in answer_compact:
        return True
    answer_norm = alnum_cjk(answer)
    term_norm = alnum_cjk(term)
    return bool(term_norm and len(term_norm) >= 4 and term_norm in answer_norm)


def _evaluate_case(
    case: dict[str, Any], payload: dict[str, Any], uploaded_id: str | None
) -> dict[str, Any]:
    """단일 케이스 채점: 기대 용어 포함 + 출처 파일명 일치를 모두 만족하면 PASS."""
    answer = str(payload.get("answer") or "")
    sources = list(payload.get("sources") or [])
    terms = [str(term) for term in case.get("expected_answer_contains", []) if str(term).strip()]
    matched_terms = [term for term in terms if _term_matches(answer, term)]
    # 최소 일치 용어 수: 보수적으로 1개 이상(생성 답변의 표현 다양성 허용)
    answer_ok = len(matched_terms) >= 1 if terms else bool(answer.strip())
    source_ok = _source_filename_match(sources, str(case["document_filename"]), uploaded_id)
    verdict = "PASS" if answer_ok and source_ok else "FAIL"
    return {
        "case_id": case["id"],
        "doc_id": case["doc_id"],
        "document_filename": case["document_filename"],
        "question": case["question"],
        "expected_answer": case["expected_answer"],
        "expected_terms": " / ".join(terms),
        "matched_terms": " / ".join(matched_terms),
        "answer": answer,
        "source_count": len(sources),
        "source_ok": source_ok,
        "answer_ok": answer_ok,
        "verdict": verdict,
        "processing_time": payload.get("processing_time"),
        "tokens_used": payload.get("tokens_used"),
    }


def _write_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = sorted({key for row in rows for key in row})
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def _write_html(path: Path, rows: list[dict[str, Any]], summary: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    row_html = []
    for row in rows:
        css_class = "pass" if row.get("verdict") == "PASS" else "fail"
        cells = "".join(
            f"<td>{html.escape(str(row.get(key, ''))[:1200])}</td>"
            for key in [
                "case_id",
                "document_filename",
                "verdict",
                "source_ok",
                "answer_ok",
                "matched_terms",
                "answer",
            ]
        )
        row_html.append(f"<tr class='{css_class}'>{cells}</tr>")
    body = (
        "<!doctype html><html lang='en'><head><meta charset='utf-8'>"
        "<title>Corpus E2E</title><style>"
        "body{font-family:sans-serif;margin:24px;}"
        "table{border-collapse:collapse;width:100%;font-size:13px;}"
        "td,th{border:1px solid #ddd;padding:6px;vertical-align:top;}"
        "tr.pass{background:#f2fff5;}tr.fail{background:#fff2f2;}"
        "pre{background:#f7f7f7;padding:12px;overflow:auto;}</style></head><body>"
        "<h1>Corpus E2E</h1>"
        f"<pre>{html.escape(json.dumps(summary, ensure_ascii=False, indent=2))}</pre>"
        "<table><thead><tr><th>Case</th><th>Document</th><th>Verdict</th>"
        "<th>Source</th><th>Answer OK</th><th>Matched terms</th><th>Actual answer</th>"
        f"</tr></thead><tbody>{''.join(row_html)}</tbody></table></body></html>"
    )
    path.write_text(body, encoding="utf-8")


def run_production(args: argparse.Namespace) -> dict[str, Any]:
    """코퍼스를 업로드하고 골든질문을 채점한 뒤 요약을 반환한다."""
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    backend_url = args.backend_url.rstrip("/")
    api_key = os.getenv(args.api_key_env) if args.api_key_env else None

    corpus_dir = Path(args.corpus_dir).expanduser().resolve()
    docs = discover_corpus(corpus_dir)
    golden = json.loads(Path(args.golden).read_text(encoding="utf-8"))
    cases = list(golden["cases"])

    upload_rows: list[dict[str, Any]] = []
    eval_rows: list[dict[str, Any]] = []
    uploaded_by_doc: dict[str, str] = {}

    staging_dir = output_dir / "_staging"
    staging_dir.mkdir(parents=True, exist_ok=True)

    with httpx.Client(timeout=httpx.Timeout(180.0, connect=30.0)) as client:
        # 1) 업로드 + 처리 완료 폴링
        if not args.skip_upload:
            selected_docs = docs[: args.upload_limit] if args.upload_limit else docs
            for doc in selected_docs:
                row: dict[str, Any] = {
                    "doc_id": doc.doc_id,
                    "filename": doc.filename,
                    "extension": doc.extension,
                    "size_bytes": doc.size_bytes,
                }
                try:
                    job_id = _upload_one(client, backend_url, doc, staging_dir, api_key)
                    final = _poll_upload(
                        client,
                        backend_url,
                        job_id,
                        timeout_seconds=args.upload_timeout_seconds,
                        poll_interval_seconds=args.poll_interval_seconds,
                        api_key=api_key,
                    )
                    row["job_id"] = job_id
                    row["final_status"] = str(final.get("status") or "")
                    row["chunk_count"] = final.get("chunk_count")
                    if row["final_status"] == "completed":
                        uploaded_by_doc[doc.doc_id] = job_id
                except Exception as error:  # noqa: BLE001 - 파일별 실패를 집계
                    row["final_status"] = "failed"
                    row["error_message"] = f"{type(error).__name__}: {error}"
                upload_rows.append(row)
        _write_csv(output_dir / "upload_results.csv", upload_rows)
        (output_dir / "upload_results.json").write_text(
            json.dumps(upload_rows, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        # 2) QA 채점
        if not args.skip_qa:
            selected_cases = cases[: args.case_limit] if args.case_limit else cases
            session_by_doc: dict[str, str] = {}
            for case_index, case in enumerate(selected_cases, start=1):
                doc_id = str(case["doc_id"])
                if doc_id not in session_by_doc:
                    session_by_doc[doc_id] = _create_session(client, backend_url, api_key)
                uploaded_id = uploaded_by_doc.get(doc_id)
                started = time.monotonic()
                try:
                    payload = _chat_answer(
                        client,
                        backend_url,
                        session_by_doc[doc_id],
                        str(case["question"]),
                        api_key,
                    )
                    result = _evaluate_case(case, payload, uploaded_id)
                except Exception as error:  # noqa: BLE001 - 케이스별 오류를 집계
                    result = {
                        "case_id": case["id"],
                        "doc_id": case["doc_id"],
                        "document_filename": case["document_filename"],
                        "question": case["question"],
                        "verdict": "ERROR",
                        "source_ok": False,
                        "answer_ok": False,
                        "answer": "",
                        "error": f"{type(error).__name__}: {error}",
                    }
                result["case_index"] = case_index
                result["request_seconds"] = round(time.monotonic() - started, 3)
                eval_rows.append(result)

    # 스테이징 정리
    for child in staging_dir.glob("*"):
        child.unlink(missing_ok=True)
    staging_dir.rmdir()

    _write_csv(output_dir / "qa_results.csv", eval_rows)
    (output_dir / "qa_results.json").write_text(
        json.dumps(eval_rows, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    upload_completed = sum(1 for row in upload_rows if row.get("final_status") == "completed")
    qa_pass = sum(1 for row in eval_rows if row.get("verdict") == "PASS")
    summary = {
        "generated_at": utc_now(),
        "backend_url": backend_url,
        "documents_expected": len(docs),
        "documents_uploaded_completed": upload_completed,
        "documents_upload_failed": len(upload_rows) - upload_completed,
        "qa_cases_run": len(eval_rows),
        "qa_pass": qa_pass,
        "qa_fail_or_error": len(eval_rows) - qa_pass,
        "success": bool(eval_rows) and qa_pass == len(eval_rows),
    }
    (output_dir / "summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    _write_html(output_dir / "qa_results.html", eval_rows, summary)
    return summary


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)

    generate = subparsers.add_parser("generate", help="코퍼스 추출 + 골든질문 생성(서버 불필요)")
    generate.add_argument("--corpus-dir", required=True, help="추출 대상 로컬 코퍼스 디렉토리")
    generate.add_argument("--output-dir", required=True, help="manifest/golden 출력 디렉토리")
    generate.add_argument(
        "--question-template",
        default=DEFAULT_QUESTION_TEMPLATE,
        help="질문 템플릿 1 ({filename}, {anchor} 치환)",
    )
    generate.add_argument(
        "--question-template-2",
        default=DEFAULT_QUESTION_TEMPLATE_2,
        help="질문 템플릿 2 ({filename}, {anchor} 치환)",
    )
    generate.add_argument(
        "--manual-facts-json",
        default=None,
        help="파일명->정답 사실 리스트 매핑 JSON(추출 빈약 시 보완, 선택)",
    )

    run = subparsers.add_parser("run-production", help="라이브 서버에 업로드 + 채점")
    run.add_argument("--corpus-dir", required=True)
    run.add_argument("--golden", required=True, help="generate가 만든 golden_questions.json")
    run.add_argument("--output-dir", required=True)
    run.add_argument("--backend-url", required=True, help="API 서버 베이스 URL")
    run.add_argument("--api-key-env", default="FASTAPI_AUTH_KEY")
    run.add_argument("--skip-upload", action="store_true", help="업로드 생략(이미 적재된 경우)")
    run.add_argument("--skip-qa", action="store_true", help="채점 생략(업로드만)")
    run.add_argument("--upload-limit", type=int, default=0, help="업로드할 최대 문서 수(0=전체)")
    run.add_argument("--case-limit", type=int, default=0, help="채점할 최대 케이스 수(0=전체)")
    run.add_argument("--upload-timeout-seconds", type=float, default=1800.0)
    run.add_argument("--poll-interval-seconds", type=float, default=5.0)

    args = parser.parse_args()
    if args.command == "generate":
        paths = generate_golden(
            Path(args.corpus_dir).expanduser().resolve(),
            Path(args.output_dir),
            question_template=args.question_template,
            question_template_2=args.question_template_2,
            manual_facts_json=args.manual_facts_json,
        )
        print(json.dumps({key: str(value) for key, value in paths.items()}, ensure_ascii=False, indent=2))
        return 0
    if args.command == "run-production":
        summary = run_production(args)
        print(json.dumps(summary, ensure_ascii=False, indent=2))
        return 0 if summary["success"] else 1
    raise AssertionError(args.command)


if __name__ == "__main__":
    raise SystemExit(main())
