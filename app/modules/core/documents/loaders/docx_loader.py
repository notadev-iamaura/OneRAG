"""
DOCX Document Loader
Word 문서 파일 로딩 전략 구현

차용/개선 사항(JapanRAG 백포트, #26):
- 레거시 .doc 지원: LibreOffice headless 변환(soffice --convert-to docx) 후 기존
  DOCX 경로를 재사용한다.
- graceful-optional: soffice/libreoffice 바이너리 부재 시 명확한 ValueError로 안내하고
  변환을 수행하지 않는다(필수 의존성 추가 금지 — OneRAG Docker/extras에 강제하지 않음).
- 범용화: JapanRAG의 macOS 전용 하드코딩 경로(/opt/homebrew/bin/soffice) 대신
  환경변수(ONERAG_SOFFICE_PATH)로 외부화하고 PATH 탐지를 우선한다.
"""

import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document

from .....lib.logger import get_logger
from .base import DocumentLoaderStrategy

logger = get_logger(__name__)

# 레거시 .doc 변환용 LibreOffice 바이너리 경로 오버라이드(환경변수). 기본 None.
SOFFICE_PATH_ENV = "ONERAG_SOFFICE_PATH"
# 변환 타임아웃(초). 무한 대기 방지.
_DOC_CONVERSION_TIMEOUT_SECONDS = 120


def _resolve_soffice_override() -> str | None:
    """환경변수로 지정된 soffice 바이너리 경로를 반환한다(실존 시에만)."""
    override = os.getenv(SOFFICE_PATH_ENV)
    if override and Path(override).exists():
        return override
    return None


class DOCXLoader(DocumentLoaderStrategy):
    """Word 문서 로더(.docx 및 레거시 .doc 지원)"""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".DOCX", ".doc", ".DOC"]

    async def load(self, file_path: Path) -> list[Document]:
        """DOC/DOCX 파일 로드"""
        if file_path.suffix.lower() == ".doc":
            return self._load_legacy_doc(file_path)
        return self._load_docx(file_path)

    def _load_docx(self, file_path: Path) -> list[Document]:
        """DOCX 파일 로드"""
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                logger.warning(f"Empty DOCX file: {file_path.name}")
                return []
            content = "\n".join(paragraphs)
            documents = [Document(page_content=content, metadata={})]
            logger.info(f"DOCX loaded: {len(paragraphs)} paragraphs from {file_path.name}")
            return documents
        except Exception as e:
            logger.error(f"DOCX loading failed for {file_path}: {e}")
            raise ValueError(f"Failed to load DOCX file: {e}") from e

    def _resolve_soffice(self) -> str:
        """LibreOffice/soffice 바이너리를 탐지한다(환경변수 → PATH 순).

        Raises:
            ValueError: 바이너리를 찾지 못한 경우(필수 의존성 미추가 안내).
        """
        soffice = (
            _resolve_soffice_override()
            or shutil.which("soffice")
            or shutil.which("libreoffice")
        )
        if not soffice:
            raise ValueError(
                "Legacy .doc files require LibreOffice/soffice for conversion, but it was "
                f"not found. Install LibreOffice or set {SOFFICE_PATH_ENV} to the soffice path."
            )
        return soffice

    def _load_legacy_doc(self, file_path: Path) -> list[Document]:
        """LibreOffice로 구형 .doc를 .docx로 변환한 뒤 DOCX 경로를 재사용한다(#26).

        graceful-optional: soffice 부재 시 변환을 수행하지 않고 명확한 ValueError를 던진다.
        임시 디렉토리는 컨텍스트 매니저로 항상 정리되며, 변환은 timeout으로 보호된다.
        """
        soffice = self._resolve_soffice()
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                subprocess.run(
                    [
                        soffice,
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        tmpdir,
                        str(file_path),
                    ],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=_DOC_CONVERSION_TIMEOUT_SECONDS,
                )
                converted_path = Path(tmpdir) / f"{file_path.stem}.docx"
                if not converted_path.exists():
                    candidates = list(Path(tmpdir).glob("*.docx"))
                    if not candidates:
                        raise ValueError("LibreOffice did not produce a .docx file.")
                    converted_path = candidates[0]
                return self._load_docx(converted_path)
        except subprocess.TimeoutExpired as e:
            logger.error(f"DOC conversion timed out for {file_path}")
            raise ValueError(
                f"DOC conversion timed out after {_DOC_CONVERSION_TIMEOUT_SECONDS} seconds"
            ) from e
        except subprocess.CalledProcessError as e:
            logger.error(f"DOC conversion failed for {file_path}: {e.stderr or e.stdout}")
            raise ValueError(f"DOC conversion failed: {e.stderr or e.stdout}") from e
